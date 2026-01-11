# =============================================================================
# Resume Parser - Text Extraction from Resume Files
# =============================================================================

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class ResumeParserError(Exception):
    """Base exception for resume parsing errors."""
    pass


class ResumeParser:
    """
    Extract text from resume files.
    
    Supports:
    - PDF (via pdfplumber)
    - DOCX (via python-docx)
    - TXT (direct read)
    - RTF (via striprtf)
    """
    
    # Supported MIME types
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "text/plain": "txt",
        "text/rtf": "rtf",
        "application/rtf": "rtf",
    }
    
    def extract_text(
        self,
        file_bytes: bytes,
        content_type: Optional[str] = None,
        filename: Optional[str] = None,
    ) -> str:
        """
        Extract text from a resume file.
        
        Args:
            file_bytes: Raw file content
            content_type: MIME type of the file
            filename: Original filename (for extension detection)
        
        Returns:
            Extracted text content
        
        Raises:
            ResumeParserError: If extraction fails
        """
        # Determine file type
        file_type = self._detect_file_type(content_type, filename)
        
        if not file_type:
            raise ResumeParserError(
                f"Unsupported file type: {content_type or 'unknown'}"
            )
        
        # Extract based on type
        try:
            if file_type == "pdf":
                return self._extract_pdf(file_bytes)
            elif file_type == "docx":
                return self._extract_docx(file_bytes)
            elif file_type == "txt":
                return self._extract_txt(file_bytes)
            elif file_type == "rtf":
                return self._extract_rtf(file_bytes)
            elif file_type == "doc":
                # DOC format is legacy and harder to parse
                # Try as RTF first, then fall back
                try:
                    return self._extract_rtf(file_bytes)
                except Exception:
                    raise ResumeParserError(
                        "Legacy .doc format not fully supported. "
                        "Please convert to .docx or .pdf"
                    )
            else:
                raise ResumeParserError(f"Unsupported file type: {file_type}")
                
        except ResumeParserError:
            raise
        except Exception as e:
            logger.exception(f"Failed to extract text from {file_type}")
            raise ResumeParserError(f"Extraction failed: {str(e)}")
    
    def _detect_file_type(
        self,
        content_type: Optional[str],
        filename: Optional[str],
    ) -> Optional[str]:
        """Detect file type from MIME type or filename."""
        # Try MIME type first
        if content_type and content_type in self.SUPPORTED_TYPES:
            return self.SUPPORTED_TYPES[content_type]
        
        # Fall back to filename extension
        if filename:
            ext = filename.lower().split(".")[-1]
            extension_map = {
                "pdf": "pdf",
                "docx": "docx",
                "doc": "doc",
                "txt": "txt",
                "rtf": "rtf",
            }
            if ext in extension_map:
                return extension_map[ext]
        
        return None
    
    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        import pdfplumber
        
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        if not text_parts:
            logger.warning("PDF appears to be empty or image-based")
            return ""
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    
    def _extract_txt(self, file_bytes: bytes) -> str:
        """Extract text from plain text file."""
        # Try common encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with replacement
        return file_bytes.decode("utf-8", errors="replace")
    
    def _extract_rtf(self, file_bytes: bytes) -> str:
        """Extract text from RTF using striprtf."""
        from striprtf.striprtf import rtf_to_text
        
        rtf_content = file_bytes.decode("utf-8", errors="replace")
        return rtf_to_text(rtf_content)


# Singleton instance
resume_parser = ResumeParser()
