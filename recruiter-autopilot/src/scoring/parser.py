"""Resume text extraction from PDF and DOCX files."""

import io
import re
from pathlib import Path

from src.utils.logging import get_logger
from src.utils.security import sanitize_resume_content

logger = get_logger(__name__)


class ResumeParseError(Exception):
    """Raised when resume parsing fails."""
    pass


class ResumeParser:
    """
    Extracts text content from resume files.

    Supports:
    - PDF files (using pypdf and pdfplumber)
    - DOCX files (using python-docx)

    Security:
    - Sanitizes extracted content to prevent prompt injection
    - Limits content size to prevent memory issues
    """

    MAX_CONTENT_SIZE = 100_000  # 100KB of text

    def __init__(self):
        pass

    def parse(self, content: bytes, filename: str) -> str:
        """
        Parse resume content from file bytes.

        Args:
            content: Raw file bytes
            filename: Original filename (used to determine type)

        Returns:
            Extracted text content

        Raises:
            ResumeParseError: If parsing fails
        """
        extension = Path(filename).suffix.lower()

        try:
            if extension == ".pdf":
                text = self._parse_pdf(content)
            elif extension in (".docx", ".doc"):
                text = self._parse_docx(content)
            elif extension == ".txt":
                text = self._parse_text(content)
            else:
                raise ResumeParseError(f"Unsupported file type: {extension}")

            # Sanitize to prevent prompt injection
            text = sanitize_resume_content(text)

            # Limit size
            if len(text) > self.MAX_CONTENT_SIZE:
                logger.warning(
                    "Resume content truncated",
                    original_size=len(text),
                    max_size=self.MAX_CONTENT_SIZE,
                )
                text = text[: self.MAX_CONTENT_SIZE]

            # Clean up whitespace
            text = self._clean_text(text)

            logger.debug(
                "Resume parsed successfully",
                filename=filename,
                content_length=len(text),
            )

            return text

        except Exception as e:
            logger.error(
                "Resume parsing failed",
                filename=filename,
                error=str(e),
            )
            raise ResumeParseError(f"Failed to parse {filename}: {str(e)}")

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF using multiple strategies."""
        text_parts = []

        # Try pdfplumber first (better for complex layouts)
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            if text_parts:
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.debug("pdfplumber extraction failed, trying pypdf", error=str(e))

        # Fallback to pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if text_parts:
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.debug("pypdf extraction failed", error=str(e))

        if not text_parts:
            raise ResumeParseError("Could not extract text from PDF")

        return "\n\n".join(text_parts)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX file."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        if not text_parts:
            raise ResumeParseError("Could not extract text from DOCX")

        return "\n".join(text_parts)

    def _parse_text(self, content: bytes) -> str:
        """Parse plain text file."""
        # Try common encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        raise ResumeParseError("Could not decode text file")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Replace multiple newlines with double newline
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Replace multiple spaces with single space
        text = re.sub(r" {2,}", " ", text)

        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        # Remove empty lines at start/end
        text = text.strip()

        return text


class MockResumeParser(ResumeParser):
    """Mock parser for testing that returns synthetic resume content."""

    def parse(self, content: bytes, filename: str) -> str:
        """Return mock resume content for testing."""
        return """
JANE DOE
Senior Software Engineer
jane.doe@email.com | (555) 123-4567 | San Francisco, CA
linkedin.com/in/janedoe | github.com/janedoe

SUMMARY
Experienced software engineer with 7+ years building scalable backend systems.
Strong expertise in Python, distributed systems, and cloud infrastructure.
Led teams of 5-8 engineers on critical platform initiatives.

EXPERIENCE

Senior Software Engineer | TechCorp Inc. | 2020 - Present
- Designed and implemented microservices architecture serving 10M+ requests/day
- Led migration from monolith to microservices, reducing deployment time by 80%
- Mentored 5 junior engineers, conducting code reviews and technical guidance
- Technologies: Python, Go, Kubernetes, PostgreSQL, Redis, AWS

Software Engineer | StartupXYZ | 2017 - 2020
- Built real-time data pipeline processing 1TB+ daily using Apache Kafka
- Developed REST APIs handling 50K concurrent users
- Improved system reliability from 99.5% to 99.99% uptime
- Technologies: Python, Django, Kafka, Docker, GCP

Junior Developer | WebAgency LLC | 2015 - 2017
- Developed web applications for 20+ clients
- Created automated testing suites reducing bugs by 40%
- Technologies: Python, JavaScript, PostgreSQL

EDUCATION

M.S. Computer Science | Stanford University | 2015
B.S. Computer Science | UC Berkeley | 2013

SKILLS

Languages: Python (expert), Go, JavaScript, SQL
Frameworks: FastAPI, Django, React
Infrastructure: AWS, GCP, Kubernetes, Docker, Terraform
Databases: PostgreSQL, Redis, MongoDB, Elasticsearch
Other: System Design, API Design, Agile/Scrum, Technical Leadership

CERTIFICATIONS
- AWS Solutions Architect Professional
- Google Cloud Professional Data Engineer
"""
